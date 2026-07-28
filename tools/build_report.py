from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "report"
ASSETS = OUT / "assets"
DOCX_PATH = OUT / "RentalApp_40535392_Report_Final.docx"
REPOSITORY_URL = "https://github.com/justinwylie033/RentalApp_40535392"
PULL_REQUEST_URL = f"{REPOSITORY_URL}/pull/1"
CI_RUN_URL = f"{REPOSITORY_URL}/actions/runs/30217148542"
CI_COMMIT = "8a913ee"
FINAL_TESTS = 89
FINAL_LINE_COVERAGE = 87.4
FINAL_BRANCH_COVERAGE = 65.7
FINAL_METHOD_COVERAGE = 87.6
FINAL_COMMITS = 18
SUBMISSION_DATE = "27 July 2026"
CI_RUN_DATE = "26 July 2026"

# Selected design system: standard_business_brief.
# Named coursework overrides: Calibri 12 pt body and single line spacing to meet
# the assessment brief; the first page uses one restrained editorial cover.
NAVY = "17324D"
BLUE = "2E74B5"
TEAL = "1D7A8C"
SKY = "DDEBF7"
LIGHT = "F2F4F7"
PALE = "F7FAFC"
INK = "1E293B"
MUTED = "52606D"
GREEN = "19734C"
AMBER = "9A6700"
RED = "A61B1B"
WHITE = "FFFFFF"
GRID = "C8D2DC"
CONTENT_DXA = 9360


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def font(size: int, bold: bool = False):
    try:
        return ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", size=size)
    except OSError:
        return ImageFont.load_default()


def font_bold(size: int):
    try:
        return ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", size=size)
    except OSError:
        return font(size, True)


def rounded(draw: ImageDraw.ImageDraw, box, fill, outline=GRID, radius=22, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, start, end, color=NAVY, width=5, head=15):
    draw.line([start, end], fill=f"#{color}", width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    left = (x2 - ux * head + px * head * 0.55, y2 - uy * head + py * head * 0.55)
    right = (x2 - ux * head - px * head * 0.55, y2 - uy * head - py * head * 0.55)
    draw.polygon([end, left, right], fill=f"#{color}")


def centered(draw, box, text, face, fill=INK):
    x1, y1, x2, y2 = box
    bb = draw.multiline_textbbox((0, 0), text, font=face, align="center", spacing=7)
    w, h = bb[2] - bb[0], bb[3] - bb[1]
    draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), text, font=face,
                        fill=f"#{fill}", align="center", spacing=7)


def save_component_diagram(path: Path):
    im = Image.new("RGB", (1800, 900), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 52), "RentalApp component architecture", font=font_bold(44), fill=f"#{NAVY}")
    d.text((80, 112), "Dependencies flow through interfaces; business rules remain at the API boundary.",
           font=font(23), fill=f"#{MUTED}")
    boxes = [
        (70, 280, 320, 540, "MAUI Views\nXAML + Shell", SKY),
        (360, 280, 610, 540, "Application\nMVVM + services", PALE),
        (650, 280, 900, 540, "ASP.NET Core\nMinimal API + JWT", SKY),
        (940, 280, 1190, 540, "Domain services\nState machine", PALE),
        (1230, 280, 1480, 540, "EF Core\nRepositories", SKY),
        (1520, 280, 1770, 540, "PostgreSQL 16\nPostGIS", "E5F5F2"),
    ]
    for x1, y1, x2, y2, label, fill in boxes:
        rounded(d, (x1, y1, x2, y2), f"#{fill}", f"#{GRID}")
        centered(d, (x1, y1, x2, y2), label, font_bold(26))
    for i in range(len(boxes) - 1):
        arrow(d, (boxes[i][2] + 8, 410), (boxes[i + 1][0] - 8, 410), BLUE, 5, 16)
    tags = [
        (90, 650, "Bindings and commands"),
        (390, 650, "Typed HTTP + DTOs"),
        (690, 650, "Authorisation"),
        (970, 650, "Workflow rules"),
        (1260, 650, "Persistence"),
        (1540, 650, "Spatial query"),
    ]
    for x, y, t in tags:
        d.text((x, y), t, font=font(19), fill=f"#{MUTED}")
    im.save(path, quality=95)


def save_er_diagram(path: Path):
    im = Image.new("RGB", (1800, 1100), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "Database schema", font=font_bold(44), fill=f"#{NAVY}")
    d.text((70, 105), "UUID keys, explicit relationships, spatial indexing, and review integrity.",
           font=font(23), fill=f"#{MUTED}")

    entities = {
        "User": (90, 235, ["Id (PK)", "Email (unique)", "DisplayName", "PasswordHash"]),
        "Item": (680, 205, ["Id (PK)", "OwnerId (FK)", "Title", "DailyRate", "Address", "Location geography(POINT,4326)"]),
        "Rental": (680, 665, ["Id (PK)", "ItemId (FK)", "BorrowerId (FK)", "Start / End UTC", "TotalPrice", "Status"]),
        "Review": (1280, 665, ["Id (PK)", "RentalId (FK, unique)", "ItemId (FK)", "ReviewerId (FK)", "Rating 1-5", "Comment"]),
        "RefreshToken": (90, 720, ["Id (PK)", "UserId (FK)", "TokenHash (unique)", "ExpiresAtUtc", "RevokedAtUtc"]),
    }
    boxes = {}
    for name, (x, y, fields) in entities.items():
        w = 460 if name != "Item" else 560
        h = 90 + len(fields) * 48
        boxes[name] = (x, y, x + w, y + h)
    # Relationships first so boxes sit cleanly above lines.
    arrow(d, (boxes["User"][2], 370), (boxes["Item"][0], 370), TEAL, 4, 14)
    d.text((500, 330), "owns 1..*", font=font(20), fill=f"#{TEAL}")
    arrow(d, (boxes["User"][2], 480), (boxes["Rental"][0], 755), TEAL, 4, 14)
    d.text((470, 565), "borrows 1..*", font=font(20), fill=f"#{TEAL}")
    arrow(d, (boxes["User"][2], 800), (boxes["RefreshToken"][2], 800), TEAL, 4, 14)
    d.text((250, 755), "receives", font=font(20), fill=f"#{TEAL}")
    arrow(d, (boxes["Item"][0] + 280, boxes["Item"][3]), (boxes["Rental"][0] + 280, boxes["Rental"][1]), BLUE, 4, 14)
    d.text((990, 590), "booked as 1..*", font=font(20), fill=f"#{BLUE}")
    arrow(d, (boxes["Rental"][2], 805), (boxes["Review"][0], 805), BLUE, 4, 14)
    d.text((1145, 765), "produces 0..1", font=font(20), fill=f"#{BLUE}")
    for name, box in boxes.items():
        x1, y1, x2, y2 = box
        rounded(d, box, f"#{PALE}", f"#{GRID}", radius=18, width=3)
        d.rounded_rectangle((x1, y1, x2, y1 + 70), radius=18, fill=f"#{NAVY}")
        d.rectangle((x1, y1 + 35, x2, y1 + 70), fill=f"#{NAVY}")
        d.text((x1 + 24, y1 + 17), name, font=font_bold(27), fill="white")
        for i, field in enumerate(entities[name][2]):
            d.text((x1 + 24, y1 + 90 + i * 48), field, font=font(21), fill=f"#{INK}")
    d.text((1285, 1030), "GiST index: Item.Location", font=font_bold(21), fill=f"#{GREEN}")
    im.save(path, quality=95)


def save_sequence_diagram(path: Path):
    im = Image.new("RGB", (1800, 1120), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "Rental request sequence", font=font_bold(44), fill=f"#{NAVY}")
    d.text((70, 105), "The API authenticates the caller and enforces booking rules before persistence.",
           font=font(23), fill=f"#{MUTED}")
    actors = [(145, "Borrower"), (475, "MAUI App"), (805, "Rental API"), (1135, "Workflow Service"), (1510, "PostGIS DB")]
    for x, name in actors:
        rounded(d, (x - 110, 190, x + 110, 275), f"#{SKY}", f"#{GRID}", 15, 3)
        centered(d, (x - 110, 190, x + 110, 275), name, font_bold(23))
        d.line((x, 275, x, 1030), fill=f"#{GRID}", width=3)
    events = [
        (345, 145, 475, "Choose item and dates"),
        (440, 475, 805, "POST /rentals + JWT"),
        (535, 805, 1135, "RequestAsync(user, dates)"),
        (630, 1135, 1510, "Check item + date overlap"),
        (725, 1510, 1135, "No conflicting booking"),
        (820, 1135, 1510, "INSERT Requested rental"),
        (915, 1135, 805, "DTO + inclusive total"),
        (1000, 805, 475, "201 Created / confirmation"),
    ]
    for y, x1, x2, label in events:
        colour = GREEN if x2 < x1 else BLUE
        arrow(d, (x1, y), (x2, y), colour, 4, 14)
        mid = (x1 + x2) / 2
        bb = d.textbbox((0, 0), label, font=font(20))
        d.rectangle((mid - (bb[2]-bb[0])/2 - 10, y - 33, mid + (bb[2]-bb[0])/2 + 10, y - 4), fill="white")
        d.text((mid - (bb[2]-bb[0])/2, y - 31), label, font=font(20), fill=f"#{INK}")
    im.save(path, quality=95)


def save_state_diagram(path: Path):
    im = Image.new("RGB", (1800, 980), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "Rental state machine", font=font_bold(44), fill=f"#{NAVY}")
    d.text((70, 105), "State classes validate workflow transitions; the service validates the acting user's role.",
           font=font(23), fill=f"#{MUTED}")
    nodes = {
        "Requested": (245, 365),
        "Approved": (600, 365),
        "Rejected": (600, 700),
        "Cancelled": (955, 700),
        "Out for rent": (955, 365),
        "Overdue": (1285, 700),
        "Returned": (1335, 365),
        "Completed": (1655, 365),
    }
    def box(name):
        x, y = nodes[name]
        return (x - 125, y - 55, x + 125, y + 55)
    edges = [
        ("Requested", "Approved"),
        ("Requested", "Rejected"),
        ("Requested", "Cancelled"),
        ("Approved", "Rejected"),
        ("Approved", "Cancelled"),
        ("Approved", "Out for rent"),
        ("Out for rent", "Overdue"),
        ("Out for rent", "Returned"),
        ("Overdue", "Returned"),
        ("Returned", "Completed"),
    ]
    for a, b in edges:
        x1, y1 = nodes[a]; x2, y2 = nodes[b]
        dx, dy = x2-x1, y2-y1
        length = max((dx*dx+dy*dy)**0.5, 1)
        ux, uy = dx/length, dy/length
        arrow(d, (x1+ux*130, y1+uy*60), (x2-ux*130, y2-uy*60), TEAL if "Overdue" in (a,b) else BLUE, 4, 15)
    for name, (x, y) in nodes.items():
        fill = "FFF4D6" if name == "Overdue" else ("E6F4EA" if name == "Completed" else SKY)
        rounded(d, box(name), f"#{fill}", f"#{GRID}", 20, 3)
        centered(d, box(name), name, font_bold(23))
    d.ellipse((70, 335, 130, 395), fill=f"#{NAVY}")
    arrow(d, (130, 365), (120, 365), NAVY, 4, 14)
    d.text((75, 420), "Start", font=font(18), fill=f"#{MUTED}")
    d.text((70, 865), "Automatic transition: Out for rent -> Overdue when the due date passes.",
           font=font_bold(22), fill=f"#{GREEN}")
    im.save(path, quality=95)


def save_coverage_evidence(path: Path):
    """Create a compact, traceable rendering of the final CI coverage export."""
    im = Image.new("RGB", (1800, 900), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 48), "Verified test and coverage export", font=font_bold(44), fill=f"#{NAVY}")
    d.text(
        (70, 108),
        f"GitHub Actions run #3 | {CI_RUN_DATE} | commit 8a913ee",
        font=font(23),
        fill=f"#{MUTED}",
    )

    cards = [
        ("Tests", "89 / 89", "passed", GREEN),
        ("Line", "87.4%", "1448 / 1656", BLUE),
        ("Branch", "65.7%", "244 / 371", TEAL),
        ("Method", "87.6%", "277 / 316", NAVY),
    ]
    x_positions = [70, 500, 930, 1360]
    for x, (label, value, detail, colour) in zip(x_positions, cards):
        rounded(d, (x, 205, x + 370, 430), f"#{PALE}", f"#{colour}", 20, 4)
        d.text((x + 26, 232), label, font=font_bold(24), fill=f"#{MUTED}")
        d.text((x + 26, 282), value, font=font_bold(48), fill=f"#{colour}")
        d.text((x + 26, 365), detail, font=font(21), fill=f"#{MUTED}")

    d.text((70, 505), "Handwritten assembly line coverage", font=font_bold(27), fill=f"#{NAVY}")
    assemblies = [
        ("RentalApp.Api", 89.7, BLUE),
        ("RentalApp.Application", 81.4, TEAL),
        ("RentalApp.Database", 94.9, GREEN),
    ]
    for index, (name, percentage, colour) in enumerate(assemblies):
        y = 575 + index * 86
        d.text((70, y), name, font=font_bold(22), fill=f"#{INK}")
        d.rounded_rectangle((490, y + 2, 1615, y + 34), radius=16, fill=f"#{LIGHT}")
        bar_end = 490 + int(1125 * percentage / 100)
        d.rounded_rectangle((490, y + 2, bar_end, y + 34), radius=16, fill=f"#{colour}")
        d.text((1640, y - 1), f"{percentage:.1f}%", font=font_bold(22), fill=f"#{colour}")

    d.text(
        (70, 846),
        "Source: test-results-and-coverage artifact (tests.trx and CoverageReport/Cobertura.xml).",
        font=font(18),
        fill=f"#{MUTED}",
    )
    im.save(path, quality=95)


def save_ci_evidence(path: Path):
    """Render the API-verified successful workflow summary without imitating a screenshot."""
    im = Image.new("RGB", (1800, 900), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 48), "GitHub Actions - verified successful run", font=font_bold(44), fill=f"#{NAVY}")
    d.text(
        (70, 108),
        "Build, test and package | run #3 | pull request #1 | commit 8a913ee",
        font=font(23),
        fill=f"#{MUTED}",
    )

    jobs = [
        (
            70,
            "backend-tests",
            [
                "PostGIS 16 service healthy",
                "89 tests passed",
                "87.4% line coverage (80% gate passed)",
                "Coverage artifact uploaded",
            ],
        ),
        (
            930,
            "android-build",
            [
                ".NET 10 + MAUI Android",
                "Android API 36",
                "Signed Release APK published",
                "APK artifact uploaded",
            ],
        ),
    ]
    for x, title, lines in jobs:
        rounded(d, (x, 210, x + 800, 620), "#F7FBF8", f"#{GREEN}", 24, 5)
        d.ellipse((x + 34, 246, x + 86, 298), fill=f"#{GREEN}")
        centered(d, (x + 34, 246, x + 86, 298), "OK", font_bold(16), WHITE)
        d.text((x + 110, 245), title, font=font_bold(31), fill=f"#{GREEN}")
        d.text((x + 110, 292), "Success", font=font_bold(23), fill=f"#{GREEN}")
        for index, line in enumerate(lines):
            y = 365 + index * 58
            d.ellipse((x + 43, y + 7, x + 61, y + 25), fill=f"#{GREEN}")
            d.text((x + 82, y), line, font=font(22), fill=f"#{INK}")

    rounded(d, (70, 685, 1730, 820), f"#{PALE}", f"#{GRID}", 18, 3)
    d.text((100, 714), "Artifacts", font=font_bold(24), fill=f"#{NAVY}")
    d.text(
        (100, 758),
        "test-results-and-coverage (515,650 bytes)   |   rentalapp-android (27,798,845 bytes)",
        font=font(22),
        fill=f"#{INK}",
    )
    d.text((70, 852), "Source: GitHub Actions run 30217148542.", font=font(18), fill=f"#{MUTED}")
    im.save(path, quality=95)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Iterable[int], indent=120):
    widths = list(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_run(run, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, 9, MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.extend([r_fonts, color, underline, size])
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(88)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, 30, NAVY, True)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(28)
        r2 = p2.add_run(subtitle)
        set_run(r2, 16, TEAL, False)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text="", bold_prefix=None, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, 12, INK, True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, 12, INK, False, italic)
    else:
        r = p.add_run(text)
        set_run(r, 12, INK, False, italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    set_run(r, 12, INK)
    return p


def add_callout(doc, title, text, tone="blue"):
    colours = {
        "blue": (SKY, NAVY), "green": ("E6F4EA", GREEN),
        "amber": ("FFF4D6", AMBER), "red": ("FDE8E7", RED)
    }
    fill, colour = colours[tone]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, colour, 9)
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, 12, colour, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run(r2, 11, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run(r, font_size, NAVY, True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, font_size, INK)
    return table


def add_figure(doc, image_path, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    picture = p.add_run().add_picture(str(image_path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split(".", 1)[0])
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(4)
    c.paragraph_format.space_after = Pt(6)
    r = c.add_run(caption)
    set_run(r, 10, MUTED, False, True)


def add_code(doc, code, title=None):
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        set_run(r, 10, MUTED, True)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, GRID, 4)
    cell = table.cell(0, 0)
    shade(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(code.strip())
    set_run(r, 7.8, "263238", False, False, "Consolas")
    return table


def page_break(doc):
    doc.add_page_break()


def screenshot_grid(doc, slots, rows, columns, image_height_inches):
    expected_slots = rows * columns
    if len(slots) != expected_slots:
        raise ValueError(
            f"Screenshot grid requires {expected_slots} items, received {len(slots)}."
        )

    table = doc.add_table(rows=rows, cols=columns)
    base_width = CONTENT_DXA // columns
    widths = [base_width] * columns
    widths[-1] += CONTENT_DXA - sum(widths)
    set_table_geometry(table, widths, indent=0)
    set_table_borders(table, GRID, 8)
    for cell, slot in zip([c for row in table.rows for c in row.cells], slots):
        title, evidence, image_path = slot
        image_path = Path(image_path)
        if not image_path.exists():
            raise FileNotFoundError(
                f"Required emulator screenshot is missing: {image_path}"
            )

        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(title)
        set_run(r, 12, NAVY, True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(5)
        picture = p2.add_run().add_picture(
            str(image_path),
            height=Inches(image_height_inches),
        )
        picture._inline.docPr.set("descr", f"Running RentalApp screenshot: {title}")
        picture._inline.docPr.set("title", title)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(evidence)
        set_run(r3, 9, MUTED, False, True)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(12)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    headings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, colour, before, after) in headings.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def configure_page(doc):
    section = doc.sections[0]
    # Define both right/odd and left/even page furniture explicitly. LibreOffice
    # otherwise creates a separate left-page style during PDF conversion and
    # can drop the static header/footer text on even pages.
    doc.settings.odd_and_even_pages_header_footer = True
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    # Keep the cover visually quiet. Running coursework furniture begins on
    # page 2, while the first-page title block contains the required metadata.
    clear_paragraph(section.first_page_header.paragraphs[0])
    clear_paragraph(section.first_page_footer.paragraphs[0])

    for header in (section.header, section.even_page_header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("SET09102 COURSEWORK  |  RENTALAPP")
        set_run(r, 9, MUTED, True)
    for footer in (section.footer, section.even_page_footer):
        add_page_number(footer.paragraphs[0])


def build_document():
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "component": ASSETS / "component.png",
        "schema": ASSETS / "schema.png",
        "sequence": ASSETS / "sequence.png",
        "state": ASSETS / "state.png",
        "coverage": ASSETS / "coverage-evidence.png",
        "ci": ASSETS / "ci-evidence.png",
    }
    save_component_diagram(diagrams["component"])
    save_er_diagram(diagrams["schema"])
    save_sequence_diagram(diagrams["sequence"])
    save_state_diagram(diagrams["state"])
    save_coverage_evidence(diagrams["coverage"])
    save_ci_evidence(diagrams["ci"])
    evidence_dir = ROOT / "docs" / "evidence"
    screenshots = {
        "login": evidence_dir / "login.png",
        "marketplace": evidence_dir / "marketplace.png",
        "item_details": evidence_dir / "item-details.png",
        "list_item": evidence_dir / "list-item.png",
        "rental_workflow": evidence_dir / "rental-workflow.png",
        "navigation": evidence_dir / "navigation-sign-out.png",
        "profile": evidence_dir / "profile-sign-out.png",
        "near_me": evidence_dir / "near-me.png",
    }
    near_me_source = screenshots["near_me"]
    near_me_crop = ASSETS / "near-me-evidence.png"
    if near_me_source.exists():
        with Image.open(near_me_source) as source:
            # Crop the portrait emulator capture to the controls so it remains
            # legible inside the report's compact 2x2 evidence grid.
            source.crop((0, 0, source.width, min(620, source.height))).save(near_me_crop)

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "RentalApp - Peer-to-Peer Rental Marketplace"
    doc.core_properties.subject = "SET09102 final coursework report"
    doc.core_properties.author = "Justin Wylie"
    doc.core_properties.keywords = "SET09102, .NET MAUI, PostGIS, EF Core, xUnit"
    # Keep the document metadata consistent with the visible submission date
    # instead of retaining python-docx's historical template timestamp.
    report_timestamp = datetime(2026, 7, 27, 12, 0, 0)
    doc.core_properties.created = report_timestamp
    doc.core_properties.modified = report_timestamp

    # Page 1: minimal editorial cover. The project title and matriculation
    # number remain the dominant elements, while every other required field is
    # integrated into the main title-page composition rather than a footer.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(132)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.line_spacing = 1.2
    title.paragraph_format.left_indent = Inches(0.35)
    title.paragraph_format.right_indent = Inches(0.35)
    shade_paragraph(title, NAVY)
    title_run = title.add_run("RentalApp")
    set_run(title_run, 34, WHITE, True)
    title_run.add_break()
    subtitle_run = title.add_run("Peer-to-Peer Rental Marketplace")
    set_run(subtitle_run, 16, "BFE9EF", False)

    matriculation = doc.add_paragraph()
    matriculation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    matriculation.paragraph_format.space_before = Pt(30)
    matriculation.paragraph_format.space_after = Pt(0)
    matriculation_run = matriculation.add_run("40535392")
    set_run(matriculation_run, 19, TEAL, True)

    cover_details = doc.add_paragraph()
    cover_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_details.paragraph_format.space_before = Pt(52)
    cover_details.paragraph_format.space_after = Pt(0)
    cover_details.paragraph_format.line_spacing = 1.25
    cover_details.paragraph_format.left_indent = Inches(1.1)
    cover_details.paragraph_format.right_indent = Inches(1.1)
    shade_paragraph(cover_details, LIGHT)
    student_run = cover_details.add_run("Justin Wylie")
    set_run(student_run, 11, NAVY, True)
    student_run.add_break()
    date_run = cover_details.add_run(f"Submission date: {SUBMISSION_DATE}")
    set_run(date_run, 10, MUTED)
    date_run.add_break()
    repository_label = cover_details.add_run("GitHub repository")
    set_run(repository_label, 10, MUTED, True)
    repository_label.add_break()
    add_hyperlink(cover_details, REPOSITORY_URL, REPOSITORY_URL)

    # Page 2: summary and contents.
    page_break(doc)
    add_heading(doc, "1. Project overview", 1)
    add_body(doc, "RentalApp is a compact Android marketplace in which community members list useful items, find nearby listings, request rentals, progress a controlled rental workflow, and leave verified reviews. The implementation deliberately completes the brief's Tier 1 and Tier 2 requirements and adds the State Pattern and automatic overdue detection from Tier 3.")
    add_heading(doc, "Submission claim", 2)
    add_callout(
        doc,
        "Evidence-based position",
        "The implementation completes the core marketplace, spatial, workflow, review, "
        "architecture, testing, and delivery requirements. The final GitHub Actions run "
        "measured 87.4% line coverage, exceeding the brief's 80% distinction testing threshold. "
        "The grade remains subject to marking and oral defence.",
        "green",
    )
    add_heading(doc, "Technology stack", 2)
    add_table(doc, ["Area", "Technology", "Evidence"], [
        ("Client", ".NET 10 / .NET MAUI", "Android XAML, Shell, Secure Storage, GPS, address geocoding"),
        ("API", "ASP.NET Core 10", "Minimal API endpoints, JWT bearer authentication, Problem Details"),
        ("Data", "EF Core 10", "DbContext, migrations, repositories, Unit of Work"),
        ("Spatial", "PostgreSQL 16 + PostGIS", "Geography point, SRID 4326, GiST, radius query"),
        ("Quality", "xUnit + Coverlet", "89/89 passing tests; 87.4% line coverage"),
        ("DevOps", "Docker Compose + GitHub Actions", "Healthy local services, test/coverage job, signed APK job"),
        ("Environment", "VS Code Dev Container", ".NET 10 Noble image and C# / Docker extensions"),
    ], [1550, 2500, 5310], 9.5)
    add_heading(doc, "Report map", 2)
    add_body(doc, "Architecture 3-6 | Features 7-12 | Testing 13-14 | CI/CD 15 | Patterns 16-17 | AI use 18 | References and verification 19")

    # Pages 3-6: required architecture diagrams.
    page_break(doc)
    add_heading(doc, "2. Architecture documentation", 1)
    add_heading(doc, "2.1 Component diagram", 2)
    add_figure(doc, diagrams["component"], "Figure 1. Component boundaries and dependency direction.", 6.35)
    add_body(doc, "The mobile client never connects directly to PostgreSQL. Views bind to ViewModels, application services call the authenticated HTTP client, and the API owns security and business rules. EF Core repositories isolate persistence and spatial queries. This supports testability and a clear separation of concerns.")
    add_callout(doc, "Design rationale", "Interfaces keep Android concerns out of ViewModel tests, while the API remains the authoritative boundary even if a client bypasses the UI.", "blue")

    page_break(doc)
    add_heading(doc, "2.2 Database schema", 2)
    add_figure(doc, diagrams["schema"], "Figure 2. Core entities and integrity relationships.", 6.2)
    add_body(doc, "Item.Address stores the readable collection address; the MAUI client forward-geocodes typed addresses or reverse-geocodes device GPS. Item.Location is mapped to geography (point, 4326), giving distance in metres. A GiST index supports the nearby predicate. Reviews have a unique RentalId and rating check constraint.")

    page_break(doc)
    add_heading(doc, "2.3 Rental request sequence", 2)
    add_figure(doc, diagrams["sequence"], "Figure 3. Successful rental request from the MAUI app to PostgreSQL.", 6.2)
    add_body(doc, "RentalWorkflowService validates future dates, availability, self-rental, and inclusive overlap before calculating price and committing. The response includes the Requested status and total price. Exceptions are converted to consistent API errors and displayed by the ViewModel.")

    page_break(doc)
    add_heading(doc, "2.4 Rental state diagram", 2)
    add_figure(doc, diagrams["state"], "Figure 4. Permitted workflow transitions implemented by state classes.", 6.25)
    add_body(doc, "Each of the eight state classes declares its permitted successors. Relationship authorisation is intentionally separate: listing creators respond and confirm, while requesters cancel eligible requests and confirm return. The background overdue worker checks expired out-for-rent records and applies the same state machine.")

    # Pages 7-8: checklist.
    page_break(doc)
    add_heading(doc, "3. Feature implementation checklist", 1)
    feature_rows = [
        ("YES", "Registration and sign-in", "Password hashing, JWT/refresh tokens, rate limiting, retry handling"),
        ("YES", "Authenticated client", "Bearer injection, token rotation, Secure Storage, visible sign-out"),
        ("YES", "Item catalogue", "Create, browse, details, creator edit, availability, My Listings"),
        ("YES", "Marketplace discovery", "Search, sort, category filter, paging, review-count display"),
        ("YES", "Address entry", "Forward geocoding, reverse GPS lookup, readable collection address"),
        ("YES", "Nearby discovery", "Address or GPS position, category/radius filters, distance result"),
        ("YES", "PostGIS", "Geography point, GiST index, IsWithinDistance query"),
        ("YES", "Rental request", "Future dates, inclusive price, availability calendar, overlap prevention"),
        ("YES", "Rental workflow", "Relationship-aware actions, cancellation, incoming/outgoing lists"),
        ("YES", "Verified reviews", "Requester only, completed rental only, one review per rental"),
        ("YES", "MVVM", "Observable ViewModels, RelayCommand, binding-first XAML"),
        ("YES", "Repository / Unit of Work", "Generic plus item, rental, and review repositories"),
        ("YES", "Service Layer", "Client services and API business services"),
        ("YES", "State Pattern", "Eight state classes including Cancelled, with transition tests"),
        ("YES", "Overdue detection", "Hosted worker moves expired records to Overdue"),
        ("YES", "Security / tracing", "Headers, correlation IDs, request logging, production secret inputs"),
        ("YES", "Testing", "89/89 tests; 87.4% line coverage; HTTP/PostGIS integration"),
        ("YES", "CI/CD", "Green PostGIS test/coverage job and signed Android APK job"),
        ("YES", "Docker / Dev Container", "Compose health checks and .NET 10 Noble environment"),
        ("YES", "Accessibility", "Semantic descriptions on important MAUI controls"),
        ("NO", "MediatR / CQRS Lite", "Optional; excluded to keep the system compact and explainable"),
        ("NO", "SonarCloud", "Optional; coverage artifacts are implemented instead"),
    ]
    add_table(doc, ["Done", "Requirement", "Implementation evidence"], feature_rows, [700, 2600, 6060], 8.3)
    add_body(doc, "Scope decision: MediatR and SonarCloud are optional in the brief. The implementation prioritises a defensible State Pattern, real PostGIS testing, Android workflow, security controls, and reproducible delivery.", after=0)

    # Pages 8-9: feature evidence.
    page_break(doc)
    add_heading(doc, "3.1 User-facing feature evidence", 2)
    add_body(doc, "Eight genuine Android emulator captures show the principal user-facing workflows. Each claim is also traced to its XAML, ViewModel, service, and executable test evidence, so the report combines visible proof with maintainable-code evidence.")
    add_table(doc, ["Feature", "Visible behaviour", "Traceable evidence"], [
        (
            "Login / registration",
            "One account can list and request items; validation, bounded retry, and sign-out are available.",
            "Views/LoginPage.xaml; LoginViewModel.cs; RentalApiTests; LoginViewModelTests",
        ),
        (
            "Browse marketplace",
            "Search, category filter, sort order, paging, rating, and review count.",
            "Views/ItemsListPage.xaml; ItemsListViewModel.cs; ItemsListViewModelTests",
        ),
        (
            "Item details",
            "Rate, owner, address, availability, reviews, dates, and valid relationship-aware actions.",
            "Views/ItemDetailPage.xaml; ItemDetailViewModel.cs; ItemDetailViewModelTests",
        ),
        (
            "List an item",
            "Title, description, category, daily rate, typed address lookup, or current GPS.",
            "Views/CreateItemPage.xaml; CreateItemViewModel.cs; CreateItemViewModelTests",
        ),
        (
            "My Listings",
            "Current user's listings can be reviewed and edited without a separate account role.",
            "Views/MyListingsPage.xaml; MyListingsViewModel.cs; MyListingsViewModelTests",
        ),
        (
            "Accessibility",
            "Important inputs, actions, and status elements expose semantic descriptions.",
            "MAUI XAML SemanticProperties; CI Android build",
        ),
    ], [1800, 3600, 3960], 8.7)
    add_callout(doc, "Security boundary", "The MAUI client hides invalid actions for usability, but ASP.NET Core services repeat ownership, requester, workflow, and validation checks. A modified client therefore cannot bypass the business rules.", "blue")

    page_break(doc)
    add_heading(doc, "3.2 Spatial, workflow, and review evidence", 2)
    if near_me_crop.exists():
        add_figure(doc, near_me_crop, "Figure 5. Real Android emulator capture of the PostGIS-powered Near me screen.", 3.15)
    add_table(doc, ["Feature", "Implementation and test evidence"], [
        (
            "Near me",
            "NearbyItemsPage.xaml -> NearbyItemsViewModel -> ItemService -> GET /items/nearby -> LocationService/ItemRepository -> PostGIS. LocationServiceTests, NearbyItemsViewModelTests, and ItemRepositoryTests verify the path.",
        ),
        (
            "Rental workflow",
            "RentalsPage.xaml presents only valid creator/requester actions. RentalWorkflowService and eight State classes enforce request, approve, reject, cancel, start, overdue, return, and completion rules. Fourteen state-machine cases and eight workflow-service cases cover them.",
        ),
        (
            "Verified reviews",
            "ReviewsPage.xaml offers eligible completed rentals. ReviewService rejects the wrong requester, an incomplete rental, duplicate submissions, and ratings outside 1-5. Review counts flow through listing and detail DTOs.",
        ),
    ], [2200, 7160], 9.0)
    add_callout(doc, "Demonstration route", "Use Mike to request Sarah's item; Sarah approves and starts; Mike returns it; Sarah completes it; Mike submits the verified review. This proves both sides of one transaction without separate account types.", "green")

    # Pages 10-12: eight genuine emulator captures. Two larger screenshots per
    # page keep the core journey readable; the final 2x2 grid documents spatial,
    # workflow, navigation, and profile behaviour while remaining under 20 pages.
    page_break(doc)
    add_heading(doc, "3.3 Authentication and catalogue evidence", 2)
    add_body(doc, "These genuine emulator captures show the entry point and the populated marketplace. The demo account accelerates assessment without exposing a production credential.")
    screenshot_grid(doc, [
        (
            "Login and registration",
            "Figure 6. Sign-in, account creation, validation fields, and presentation demo account.",
            screenshots["login"],
        ),
        (
            "Browse marketplace",
            "Figure 7. Listings show category, price, collection address, rating, and review count.",
            screenshots["marketplace"],
        ),
    ], rows=1, columns=2, image_height_inches=4.75)

    page_break(doc)
    add_heading(doc, "3.4 Listing detail and creation evidence", 2)
    add_body(doc, "The detail view exposes ownership-aware actions, while the creation form gathers the data required by the API and PostGIS location workflow.")
    screenshot_grid(doc, [
        (
            "Item details and creator action",
            "Figure 8. Listing detail shows price, owner, rating, address, reviews, and creator-only edit.",
            screenshots["item_details"],
        ),
        (
            "List an item and collection address",
            "Figure 9. Creation captures item data and supports typed-address lookup or device GPS.",
            screenshots["list_item"],
        ),
    ], rows=1, columns=2, image_height_inches=4.75)

    page_break(doc)
    add_heading(doc, "3.5 Workflow and account evidence", 2)
    add_body(doc, "Genuine emulator captures show rental states, nearby search, navigation, reputation, unified account access, and sign-out.")
    screenshot_grid(doc, [
        (
            "Rental workflow",
            "Figure 10. Completed and returned rental states.",
            screenshots["rental_workflow"],
        ),
        (
            "Near Me",
            "Figure 11. PostGIS radius and location controls.",
            screenshots["near_me"],
        ),
        (
            "Navigation and sign-out",
            "Figure 12. Shell navigation and persistent sign-out.",
            screenshots["navigation"],
        ),
        (
            "Profile and reputation",
            "Figure 13. Reputation, unified access, and sign-out.",
            screenshots["profile"],
        ),
    ], rows=2, columns=2, image_height_inches=2.15)

    # Pages 13-14: testing.
    page_break(doc)
    add_heading(doc, "4. Testing documentation", 1)
    add_heading(doc, "4.1 Verified result and coverage", 2)
    add_callout(doc, f"Verified GitHub Actions run - {CI_RUN_DATE}", "89 total | 89 passed | 0 failed | 0 skipped. Cobertura line coverage: 87.4%.", "green")
    add_figure(doc, diagrams["coverage"], "Figure 14. Coverage values rendered from the final CI tests.trx and Cobertura export.", 6.25)
    test_rows = [
        ("RentalApiTests", "7", "Authenticated HTTP workflows against PostgreSQL/PostGIS"),
        ("ItemRepositoryTests", "1", "Real PostgreSQL/PostGIS radius integration"),
        ("Client / API services", "11", "HTTP, refresh, authentication, item, rental, and review routes"),
        ("API application services", "30", "Validation, ownership, price, overlap, spatial, review, JWT, overdue"),
        ("RentalStateMachineTests", "14", "Valid and invalid transitions including cancellation"),
        ("ViewModels (9 classes)", "26", "Busy/error state, commands, paging, listings, reviews, profile"),
    ]
    add_table(doc, ["Test group", "Cases", "Primary concern"], test_rows, [3350, 900, 5110], 8.7)
    add_body(doc, "Total: 89 executed cases across 21 test classes. Theory and InlineData rows are counted as executed xUnit cases. The workflow enforces an 80.0% line-coverage gate.", after=0)

    page_break(doc)
    add_heading(doc, "4.2 Representative test excerpts", 2)
    add_code(doc, '''[Theory]
[InlineData(RentalStatus.Requested, RentalStatus.Approved)]
[InlineData(RentalStatus.OutForRent, RentalStatus.Overdue)]
[InlineData(RentalStatus.Returned, RentalStatus.Completed)]
public void EnsureValidTransition_AllowedTransition_DoesNotThrow(
    RentalStatus current, RentalStatus next)
{
    var exception = Record.Exception(
        () => _machine.EnsureValidTransition(current, next));
    Assert.Null(exception);
}''', "Example 1 - parameterised State Pattern test")
    add_code(doc, '''[Fact]
public async Task RequestAsync_OverlappingDates_RejectsSecondRequest()
{
    var data = await SeedUsersAndItemAsync(context);
    var start = DateTimeOffset.UtcNow.Date.AddDays(4);
    await service.RequestAsync(data.Borrower.Id,
        new(data.Item.Id, start, start.AddDays(2)));

    await Assert.ThrowsAsync<BusinessRuleException>(() =>
        service.RequestAsync(data.SecondBorrower.Id,
            new(data.Item.Id, start.AddDays(1), start.AddDays(3))));
}''', "Example 2 - business rule test")
    add_code(doc, '''[Fact]
public async Task GetNearbyAsync_ItemsInsideAndOutsideRadius_ReturnsOnlyNearbyItem()
{
    var results = await repository.GetNearbyAsync(
        55.9533, -3.1883, 2, ItemCategory.Tools);

    Assert.Contains(results, x => x.Item.Id == nearby.Id);
    Assert.DoesNotContain(results, x => x.Item.Id == distant.Id);
    Assert.All(results, x => Assert.InRange(x.DistanceMetres, 0, 2_000));
}''', "Example 3 - real PostGIS integration test")
    add_body(doc, "The tests follow Arrange-Act-Assert, use descriptive names, and combine fast in-memory/mocked tests with one real spatial database fixture where an in-memory substitute would not prove translation to PostGIS.")

    # Page 15: CI/CD.
    page_break(doc)
    add_heading(doc, "5. CI/CD evidence", 1)
    add_body(doc, "The workflow runs manually and on pushes or pull requests to main/master. The backend job starts PostgreSQL 16 with PostGIS, audits dependencies, restores, builds, executes tests, enforces the 80% line-coverage gate, creates readable reports, and uploads evidence. The Android job installs Java, .NET MAUI, and API 36, then creates an ephemeral coursework key and publishes a signed Release APK.")
    add_code(doc, '''on:
  workflow_dispatch:
  push:
    branches: [main, master]
  pull_request:
    branches: [main, master]

jobs:
  backend-tests:
    services:
      postgres:
        image: postgis/postgis:16-3.5-alpine
    steps:
      - run: dotnet test RentalApp.Test/RentalApp.Test.csproj ...
      - run: python coverage gate (minimum 80%)
      - uses: actions/upload-artifact@v7

  android-build:
    steps:
      - uses: actions/setup-dotnet@v6
        with: { workloads: maui-android }
      - run: sdkmanager platforms;android-36
      - run: dotnet publish ... -f net10.0-android
      - uses: actions/upload-artifact@v7''', ".github/workflows/build.yml (abridged)")
    add_figure(doc, diagrams["ci"], "Figure 15. API-verified summary of GitHub Actions run #3; both required jobs succeeded.", 5.8)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Evidence links: ")
    set_run(r, 10.5, INK, True)
    add_hyperlink(p, CI_RUN_URL, "successful workflow run")
    p.add_run(" | ")
    add_hyperlink(p, PULL_REQUEST_URL, "merged pull request #1")

    # Pages 16-17: design patterns. The preceding CI page is full, so the
    # keep-with-next heading starts this section on the following page.
    add_heading(doc, "6. Design patterns implementation", 1)
    add_heading(doc, "6.1 MVVM", 2)
    add_body(doc, "Views define layout and bind to observable state. ViewModels expose commands and depend on abstractions, so catalogue behaviour can be tested without a page or emulator. Generated ObservableProperty and RelayCommand members reduce repetitive notification code.")
    add_code(doc, '''public partial class ItemsListViewModel(
    IItemService items, INavigationService navigation) : ViewModelBase
{
    public ObservableCollection<ItemSummaryDto> Items { get; } = [];

    [RelayCommand]
    private Task LoadAsync() => LoadPageAsync(reset: true);

    private Task LoadPageAsync(bool reset) => RunBusyAsync(async () =>
    {
        ItemCategory? category = Enum.TryParse<ItemCategory>(
            SelectedCategory, out var parsed) ? parsed : null;
        var result = await items.GetAllAsync(
            category, SearchTerm, SelectedSortOrder, _currentPage, PageSize);
        if (reset) Items.Clear();
        foreach (var item in result.Items) Items.Add(item);
        TotalResults = result.TotalCount;
        HasMoreItems = result.HasNextPage;
    });
}''', "MVVM example - ItemsListViewModel")
    add_heading(doc, "6.2 Repository and Unit of Work", 2)
    add_body(doc, "Repositories centralise EF Core expressions and eager-loading choices. The Item repository is the only layer that knows the spatial predicate. Unit of Work provides one explicit SaveChanges boundary for service operations.")
    add_code(doc, '''var origin = new Point(longitude, latitude) { SRID = 4326 };
var radiusMetres = radiusKilometres * 1_000;

var results = await Context.Items
    .AsNoTracking()
    .Where(item => item.IsAvailable &&
        item.Location.IsWithinDistance(origin, radiusMetres))
    .Select(item => new {
        Item = item,
        DistanceMetres = item.Location.Distance(origin)
    })
    .OrderBy(result => result.DistanceMetres)
    .ToListAsync(cancellationToken);''', "Repository example - PostGIS radius query")

    page_break(doc)
    add_heading(doc, "6.3 Service Layer", 2)
    add_body(doc, "RentalWorkflowService is the authoritative use-case layer. It coordinates repositories, validates rules, calculates inclusive price, checks actor permissions, invokes the state machine, and commits. Endpoints remain thin and ViewModels cannot bypass these rules.")
    add_code(doc, '''if (item.OwnerId == borrowerId)
    throw new BusinessRuleException("You cannot rent your own item.");

if (await rentals.HasDateOverlapAsync(item.Id, start, end, cancellationToken))
    throw new BusinessRuleException("The item is already booked for those dates.");

var numberOfDays = (end.Date - start.Date).Days + 1;
var rental = new Rental {
    ItemId = item.Id,
    BorrowerId = borrowerId,
    StartDateUtc = start,
    EndDateUtc = end,
    TotalPrice = item.DailyRate * numberOfDays
};''', "Service example - validation and inclusive pricing")
    add_heading(doc, "6.4 State Pattern (advanced)", 2)
    add_body(doc, "A state object owns the allowed outgoing transitions for one RentalStatus. Adding a state or transition is localised instead of extending a large conditional. RentalWorkflowService separately checks listing-creator/requester permissions, preserving single responsibility.")
    add_code(doc, '''public abstract class RentalState : IRentalState
{
    private readonly HashSet<RentalStatus> _allowed;

    protected RentalState(params RentalStatus[] allowed)
    {
        _allowed = new HashSet<RentalStatus>(allowed);
    }

    public abstract RentalStatus Status { get; }
    public bool CanTransitionTo(RentalStatus next)
    {
        return _allowed.Contains(next);
    }
}

public sealed class OutForRentState : RentalState
{
    public OutForRentState()
        : base(RentalStatus.Overdue, RentalStatus.Returned) { }

    public override RentalStatus Status => RentalStatus.OutForRent;
}''', "Advanced pattern example - explicit state classes")

    # Page 18: AI use.
    page_break(doc)
    add_heading(doc, "7. AI tool usage", 1)
    add_body(doc, "Tools used: ChatGPT Codex for architecture, implementation support, refactoring, test design, debugging, and documentation review; Visual Studio Code C# Dev Kit for compiler diagnostics, navigation, and debugging.")
    ai_rows = [
        ("Architecture", "Separate MAUI, application, API, data, migrations, and test projects.", "Accepted for testability and security. Rejected optional MediatR/SonarCloud to control scope.", "Build, dependency, service, and integration tests."),
        ("Spatial search", "Use SRID 4326 Point, geography column, GiST, and IsWithinDistance.", "Checked X=longitude, Y=latitude; validated coordinate/radius bounds.", "Real PostGIS fixture returns near item only."),
        ("Rental workflow", "One class per state; keep role checks in the service.", "Accepted to demonstrate Open/Closed design; overlap and inclusive-price rules retained.", "Theory tests plus price, overlap, role, and overdue tests."),
        ("Android debugging", "Use Logcat to diagnose splash exit and details crash.", "Changes tied to observed stack traces: embed assemblies, correct App initialisation, harden route.", "Clean rebuild/install and focused ViewModel tests."),
        ("Address entry", "Forward-geocode typed addresses and reverse-geocode GPS.", "Store readable address plus validated PostGIS point; remove raw coordinate fields from UI.", "ViewModel and authenticated API workflow tests."),
        ("Unified accounts and review UX", "Derive actions from listing-creator/requester relationships and carry review counts through DTOs.", "Kept every API ownership and requester check; the UI is helpful but is not the security boundary.", "Final regression: 89/89 tests and 87.4% coverage."),
        ("Submission upgrades", "Add search/paging, cancellation, availability, rate limiting, readiness, tracing, security headers, and semantics.", "Implemented as 13 focused commits and reviewed through pull request #1; corrected final CI failures before merge.", "Both GitHub Actions jobs green; artifacts inspected."),
    ]
    add_table(doc, ["Interaction", "Suggestion", "Evaluation / modification", "Validation"], ai_rows, [1400, 2500, 3100, 2360], 8.0)
    add_heading(doc, "Reflection", 2)
    add_body(doc, "AI accelerated repetitive implementation and exposed alternatives, but it did not replace engineering judgement. Suggestions were narrowed to an explainable scope, compiled with warnings as errors, and tested at the appropriate boundary. Database and Android defects were diagnosed from evidence instead of accepting speculative changes. In the presentation, the student must explain every excerpt above and the trade-offs behind it.")

    # Page 19: references and final verification.
    page_break(doc)
    add_heading(doc, "8. References", 1)
    references = [
        ("Microsoft (2026), .NET MAUI documentation", "https://learn.microsoft.com/en-us/dotnet/maui/?view=net-maui-10.0"),
        ("Microsoft (2025), Spatial data in Entity Framework Core", "https://learn.microsoft.com/en-us/ef/core/modeling/spatial"),
        ("Microsoft (2025), Configure JWT bearer authentication in ASP.NET Core", "https://learn.microsoft.com/en-us/aspnet/core/security/authentication/configure-jwt-bearer-authentication?view=aspnetcore-10.0"),
        ("PostGIS Project, ST_DWithin reference", "https://postgis.net/docs/ST_DWithin.html"),
        ("NetTopologySuite project, API documentation", "https://nettopologysuite.github.io/NetTopologySuite/"),
        ("xUnit.net, Getting started with xUnit.net v3", "https://xunit.net/docs/getting-started/v3/getting-started"),
        ("Docker, Docker Compose documentation", "https://docs.docker.com/compose/"),
        ("GitHub, Building and testing .NET", "https://docs.github.com/actions/guides/building-and-testing-net"),
    ]
    for label, url in references:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(6)
        add_hyperlink(p, url, label)
    add_heading(doc, "Submission verification", 2)
    add_table(doc, ["Check", "Verified evidence"], [
        ("Public repository", f"{REPOSITORY_URL} is public; master contains at least {FINAL_COMMITS} reachable commits."),
        ("Required features", "Feature matrix, source paths, architecture diagrams, and executable tests are included."),
        ("Tests / coverage", f"{FINAL_TESTS}/{FINAL_TESTS} passed; {FINAL_LINE_COVERAGE:.1f}% line coverage; 80% CI gate passed."),
        ("CI/CD", "Run #3 succeeded for backend-tests and android-build; both evidence artifacts exist."),
        ("README / setup", "Windows startup, Docker, Android, testing, deployment, and demo workflow are documented."),
        ("Credentials", "No production secret is committed; local demo values are labelled and production Compose requires environment inputs."),
        ("Code explanation", "Architecture guide, demonstration guide, XML summaries, and presentation-point comments are included."),
        ("AI disclosure", "Tools, representative interactions, evaluation, modifications, and validation are documented above and in docs/AI_USAGE.md."),
        ("App screenshots", "Eight genuine emulator captures document authentication, browsing, details, creation, spatial search, workflow, navigation, and profile."),
        ("Report format", "Final DOCX and PDF use 12 pt body text, page numbers, all required sections, and fewer than 20 pages."),
    ], [2500, 6860], 8.2)
    add_callout(doc, "Final status", "Repository, application, architecture, tests, coverage, CI/CD, commit history, patterns, AI usage, and eight genuine emulator captures are traceable. The final Android workflow was demonstrated from the submitted source and is represented by the emulator evidence in this report.", "green")

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_document())
